"""
Transcription Service
Handles audio/video transcription using faster-whisper and GPT analysis
"""

import os
import asyncio
import tempfile
import uuid
from pathlib import Path
from typing import Dict, Optional
from datetime import datetime

import aiofiles
import httpx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.logging import logger
from app.core.config import settings


class TranscriptionService:
    """Service for transcribing audio/video and creating documents"""

    def __init__(self):
        self.upload_dir = Path("./uploads/transcriptions")
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.tasks: Dict[str, Dict] = {}

        # OpenRouter configuration
        self.openrouter_api_key = os.getenv("OPENROUTER_API_KEY")
        self.openrouter_base_url = os.getenv("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1")
        self.default_model = os.getenv("DEFAULT_MODEL", "openai/gpt-4o-mini")

    async def cleanup_old_files(self, days_old: int = 7):
        """
        Очистка только исходных файлов старше указанного количества дней
        НЕ удаляет результаты транскрипции (_transcript.docx и _audio.mp3)
        По умолчанию удаляет исходники старше 7 дней
        """
        try:
            import time

            current_time = time.time()
            cutoff_time = current_time - (days_old * 24 * 60 * 60)
            deleted_count = 0
            freed_space = 0

            # Проходим по всем файлам в директории
            for file_path in self.upload_dir.glob("*"):
                if file_path.is_file():
                    filename = file_path.name

                    # НЕ УДАЛЯЕМ результаты транскрипции - они нужны для скачивания
                    if filename.endswith('_transcript.docx') or filename.endswith('_audio.mp3'):
                        continue

                    file_mtime = file_path.stat().st_mtime

                    if file_mtime < cutoff_time:
                        file_size = file_path.stat().st_size
                        file_path.unlink()
                        deleted_count += 1
                        freed_space += file_size
                        logger.info(
                            "old_source_file_deleted",
                            file=file_path.name,
                            age_days=int((current_time - file_mtime) / 86400),
                            size_mb=round(file_size / 1024 / 1024, 2)
                        )

            if deleted_count > 0:
                logger.info(
                    "cleanup_completed",
                    deleted_files=deleted_count,
                    freed_mb=round(freed_space / 1024 / 1024, 2)
                )

            return {
                "deleted_files": deleted_count,
                "freed_mb": round(freed_space / 1024 / 1024, 2)
            }

        except Exception as e:
            logger.error(f"cleanup_error: {str(e)}")
            return {"error": str(e)}

    async def save_chunk(self, chunk: bytes, session_id: str) -> str:
        """Save audio chunk for auto-save functionality"""
        try:
            chunk_dir = self.upload_dir / "chunks" / session_id
            chunk_dir.mkdir(parents=True, exist_ok=True)

            chunk_file = chunk_dir / f"chunk_{datetime.now().timestamp()}.webm"

            async with aiofiles.open(chunk_file, "wb") as f:
                await f.write(chunk)

            logger.info(f"audio_chunk_saved session_id={session_id}, file={chunk_file}")
            return str(chunk_file)

        except Exception as e:
            logger.error("error_saving_chunk", error=str(e), session_id=session_id)
            raise

    async def process_audio(self, audio_file: Path, task_id: str) -> Dict:
        """Process audio file: transcribe and analyze"""
        try:
            # Update task status (don't overwrite if already exists from chunked upload)
            if task_id not in self.tasks:
                self.tasks[task_id] = {
                    "status": "processing",
                    "progress": 0,
                    "stage": "Начало обработки...",
                    "started_at": datetime.now().isoformat()
                }

            # Check file size
            file_size_mb = audio_file.stat().st_size / (1024 * 1024)
            logger.info(f"checking_file_size task_id={task_id}, size_mb={round(file_size_mb, 2)}")

            # ПРОСТОЕ РЕШЕНИЕ: Просто извлекаем аудиодорожку из видео (аудио намного меньше видео!)
            # Видео 218MB может содержать аудио всего 10-20MB
            self.tasks[task_id]["stage"] = "Извлечение аудио из видео..."
            self.tasks[task_id]["progress"] = 10
            logger.info(f"extracting_audio task_id={task_id}")
            audio_path = await self._extract_audio_from_video(audio_file)
            self.tasks[task_id]["progress"] = 20
            self.tasks[task_id]["stage"] = "Аудио извлечено, начало транскрипции..."

            # Step 2: Transcribe audio with faster-whisper using subprocess (50% progress)
            self.tasks[task_id]["stage"] = "Транскрипция аудио (это может занять несколько минут)..."
            logger.info(f"transcribing_audio task_id={task_id}")
            transcript = await self._transcribe_file_subprocess(audio_path)
            self.tasks[task_id]["progress"] = 60
            self.tasks[task_id]["stage"] = "Транскрипция завершена, анализ текста..."

            # Step 3: Analyze with GPT (20% progress)
            logger.info(f"analyzing_with_gpt task_id={task_id}")
            analysis = await self._analyze_with_gpt(transcript)
            self.tasks[task_id]["progress"] = 75
            self.tasks[task_id]["stage"] = "Анализ завершен, создание документа..."

            # Step 4: Create DOCX document (20% progress)
            logger.info(f"creating_document task_id={task_id}")
            docx_path = await self._create_docx(transcript, analysis, task_id)
            self.tasks[task_id]["progress"] = 90
            self.tasks[task_id]["stage"] = "Сохранение аудио файла..."

            # Step 5: Save audio file (10% progress)
            audio_output = self.upload_dir / f"{task_id}_audio.mp3"
            await self._convert_to_mp3(audio_path, audio_output)
            self.tasks[task_id]["progress"] = 100
            self.tasks[task_id]["stage"] = "Готово!"

            # Update task status
            self.tasks[task_id].update({
                "status": "completed",
                "progress": 100,
                "result": {
                    "docx_url": f"/api/v1/transcription/download/{task_id}_transcript.docx",
                    "audio_url": f"/api/v1/transcription/download/{task_id}_audio.mp3",
                    "transcript": transcript,
                    "analysis": analysis
                },
                "completed_at": datetime.now().isoformat()
            })

            # АВТООЧИСТКА: удаляем оригинальный загруженный файл для экономии места
            try:
                if audio_file.exists() and audio_file != audio_path:
                    audio_file.unlink()
                    logger.info(f"original_file_deleted file={str(audio_file)}")

                # Также удаляем временный аудио файл если он отличается от итогового
                if audio_path.exists() and audio_path != audio_output:
                    audio_path.unlink()
                    logger.info(f"temp_audio_deleted file={str(audio_path)}")
            except Exception as cleanup_error:
                logger.warning(f"cleanup_error error={str(cleanup_error)}")

            logger.info(f"transcription_completed task_id={task_id}")
            return self.tasks[task_id]

        except Exception as e:
            logger.error(f"transcription_error task_id={task_id}, error={str(e)}", exc_info=True)

            # АВТООЧИСТКА: удаляем файл даже при ошибке
            try:
                if audio_file.exists():
                    audio_file.unlink()
                    logger.info(f"failed_file_deleted file={str(audio_file)}")
            except Exception as cleanup_error:
                logger.warning(f"cleanup_error_on_failure error={str(cleanup_error)}")

            self.tasks[task_id] = {
                "status": "error",
                "progress": 0,
                "error": str(e),
                "failed_at": datetime.now().isoformat()
            }
            raise

    async def _split_video_into_parts(self, video_file: Path, task_id: str) -> list[Path]:
        """
        Split large video file into smaller parts for processing
        Target size: ~50MB per part
        """
        try:
            file_size_mb = video_file.stat().st_size / (1024 * 1024)

            # Calculate number of parts (aim for ~20MB each) - REDUCED to prevent OOM
            target_size_mb = 20  # Меньшие части = меньше памяти
            num_parts = max(2, int(file_size_mb / target_size_mb) + 1)

            logger.info(
                "splitting_video",
                task_id=task_id,
                total_size_mb=round(file_size_mb, 2),
                num_parts=num_parts
            )

            # First, get video duration using ffprobe
            probe_process = await asyncio.create_subprocess_exec(
                'ffprobe',
                '-v', 'error',
                '-show_entries', 'format=duration',
                '-of', 'default=noprint_wrappers=1:nokey=1',
                str(video_file),
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )

            stdout, stderr = await probe_process.communicate()

            if probe_process.returncode != 0:
                raise Exception(f"FFprobe error: {stderr.decode()}")

            total_duration = float(stdout.decode().strip())
            segment_duration = total_duration / num_parts

            logger.info(
                "video_duration_info",
                task_id=task_id,
                total_duration=round(total_duration, 2),
                segment_duration=round(segment_duration, 2)
            )

            # Split video into parts
            parts = []
            for i in range(num_parts):
                start_time = i * segment_duration
                part_filename = video_file.stem + f"_part{i+1}{video_file.suffix}"
                part_path = self.upload_dir / part_filename

                # Use ffmpeg to extract this segment
                split_process = await asyncio.create_subprocess_exec(
                    'ffmpeg',
                    '-i', str(video_file),
                    '-ss', str(start_time),
                    '-t', str(segment_duration),
                    '-c', 'copy',  # Copy without re-encoding (fast)
                    '-avoid_negative_ts', '1',
                    str(part_path),
                    stdout=asyncio.subprocess.PIPE,
                    stderr=asyncio.subprocess.PIPE
                )

                stdout, stderr = await split_process.communicate()

                if split_process.returncode != 0:
                    # Clean up any created parts
                    for p in parts:
                        p.unlink(missing_ok=True)
                    raise Exception(f"FFmpeg split error: {stderr.decode()}")

                parts.append(part_path)
                logger.info(
                    "part_created",
                    task_id=task_id,
                    part=i+1,
                    total=num_parts,
                    size_mb=round(part_path.stat().st_size / (1024 * 1024), 2)
                )

            logger.info(
                "video_split_complete",
                task_id=task_id,
                parts_count=len(parts)
            )

            return parts

        except Exception as e:
            logger.error(f"video_split_error task_id={task_id}, error={str(e)}")
            raise

    async def _extract_audio_from_video(self, file_path: Path) -> Path:
        """Extract audio from video file using ffmpeg"""
        try:
            # Check if file is already audio
            if file_path.suffix.lower() in ['.mp3', '.wav', '.m4a', '.flac', '.webm']:
                return file_path

            # Extract audio using ffmpeg
            output_path = file_path.with_suffix('.mp3')

            process = await asyncio.create_subprocess_exec(
                'ffmpeg',
                '-i', str(file_path),
                '-vn',  # No video
                '-acodec', 'libmp3lame',
                '-q:a', '2',  # High quality
                str(output_path),
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )

            stdout, stderr = await process.communicate()

            if process.returncode != 0:
                raise Exception(f"FFmpeg error: {stderr.decode()}")

            return output_path

        except Exception as e:
            logger.error(f"audio_extraction_error: {e}")
            raise

    async def _transcribe_file_subprocess(self, audio_path: Path) -> str:
        """Transcribe using subprocess to isolate memory - РЕШЕНИЕ ПРОБЛЕМЫ OOM"""
        try:
            import json
            logger.info(f"transcribing_via_subprocess audio_file={str(audio_path)}")

            # Создаем временный Python скрипт для транскрипции
            script = f"""
import sys
import json
from faster_whisper import WhisperModel

audio_file = "{audio_path}"
model_size = "tiny"

try:
    model = WhisperModel(model_size, device="cpu", compute_type="int8", num_workers=1)
    segments, info = model.transcribe(
        audio_file,
        beam_size=1,
        best_of=1,
        language="ru",
        vad_filter=True,
        vad_parameters=dict(min_silence_duration_ms=500),
        condition_on_previous_text=False,
        temperature=0.0
    )

    transcript_parts = []
    for segment in segments:
        transcript_parts.append(segment.text)

    transcript = " ".join(transcript_parts).strip()
    print(json.dumps({{"success": True, "transcript": transcript}}))

except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}), file=sys.stderr)
    sys.exit(1)
"""

            # Запускаем в отдельном процессе
            process = await asyncio.create_subprocess_exec(
                'python3',
                '-c',
                script,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )

            stdout, stderr = await process.communicate()

            if process.returncode != 0:
                error_msg = stderr.decode() if stderr else "Unknown error"
                logger.error("subprocess_transcription_error", error=error_msg)
                raise Exception(f"Transcription failed: {error_msg}")

            result = json.loads(stdout.decode())
            if not result.get("success"):
                raise Exception(result.get("error", "Unknown error"))

            transcript = result["transcript"]
            logger.info(f"subprocess_transcription_completed chars={len(transcript)}")
            return transcript

        except Exception as e:
            logger.error(f"subprocess_transcription_error: {e}")
            raise

    async def _transcribe_with_whisper(self, audio_path: Path) -> str:
        """Transcribe audio using faster-whisper (OPTIMIZED FOR SPEED)"""
        try:
            # Load model and transcribe (for single file processing)
            model = await self._load_whisper_model()
            transcript = await self._transcribe_with_model(audio_path, model)

            # Clean up model
            del model
            import gc
            gc.collect()

            return transcript

        except Exception as e:
            logger.error(f"whisper_transcription_error: {e}")
            raise

    async def _analyze_with_gpt(self, transcript: str) -> Dict:
        """Analyze transcript using OpenRouter GPT"""
        try:
            if not self.openrouter_api_key:
                logger.warning(f'openrouter_key_missing message="Skipping GPT analysis"')
                return {
                    "summary": "Анализ не выполнен (отсутствует API ключ)",
                    "key_points": [],
                    "action_items": []
                }

            # Prepare prompt
            prompt = f"""
Проанализируй следующую транскрипцию аудио/видео и предоставь:

1. Краткое резюме (2-3 предложения)
2. Ключевые моменты (список из 5-7 пунктов)
3. Действия и задачи, которые были упомянуты (если есть)

Транскрипция:
{transcript[:4000]}  # Limit to avoid token limits

Формат ответа должен быть структурированным и четким.
"""

            # Call OpenRouter API
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(
                    f"{self.openrouter_base_url}/chat/completions",
                    headers={
                        "Authorization": f"Bearer {self.openrouter_api_key}",
                        "Content-Type": "application/json",
                        "HTTP-Referer": "https://nikolaevcodev.ru",
                        "X-Title": "Enterprise CRM - Transcription"
                    },
                    json={
                        "model": self.default_model,
                        "messages": [
                            {
                                "role": "system",
                                "content": "Ты - эксперт по анализу текста и выделению ключевой информации. Отвечай на русском языке."
                            },
                            {
                                "role": "user",
                                "content": prompt
                            }
                        ],
                        "temperature": 0.3,
                        "max_tokens": 1000
                    }
                )

            if response.status_code != 200:
                raise Exception(f"OpenRouter API error: {response.status_code} - {response.text}")

            result = response.json()
            analysis_text = result["choices"][0]["message"]["content"]

            # Parse the analysis (simple parsing)
            analysis = {
                "raw_analysis": analysis_text,
                "summary": self._extract_section(analysis_text, "резюме", "ключевые"),
                "key_points": self._extract_list_items(analysis_text),
                "action_items": self._extract_action_items(analysis_text)
            }

            logger.info(f"gpt_analysis_completed model={self.default_model}")
            return analysis

        except Exception as e:
            logger.error(f"gpt_analysis_error: {e}")
            return {
                "summary": "Ошибка при анализе",
                "key_points": [],
                "action_items": [],
                "error": str(e)
            }

    def _extract_section(self, text: str, start_marker: str, end_marker: str) -> str:
        """Extract section from text"""
        try:
            start = text.lower().find(start_marker)
            end = text.lower().find(end_marker)
            if start != -1 and end != -1:
                section = text[start:end].strip()
                # Remove the marker itself
                lines = section.split('\n')
                if len(lines) > 1:
                    return '\n'.join(lines[1:]).strip()
            return ""
        except:
            return ""

    def _extract_list_items(self, text: str) -> list:
        """Extract list items from text"""
        items = []
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if line and (line.startswith('-') or line.startswith('•') or line[0].isdigit()):
                # Clean up the line
                clean_line = line.lstrip('-•0123456789. ').strip()
                if clean_line:
                    items.append(clean_line)
        return items[:7]  # Limit to 7 items

    def _extract_action_items(self, text: str) -> list:
        """Extract action items from text"""
        actions = []
        lines = text.split('\n')
        in_action_section = False

        for line in lines:
            line_lower = line.lower()
            if 'действи' in line_lower or 'задач' in line_lower:
                in_action_section = True
                continue

            if in_action_section:
                line = line.strip()
                if line and (line.startswith('-') or line.startswith('•') or line[0].isdigit()):
                    clean_line = line.lstrip('-•0123456789. ').strip()
                    if clean_line:
                        actions.append(clean_line)

        return actions

    async def _create_docx(self, transcript: str, analysis: Dict, task_id: str) -> Path:
        """Create DOCX document with transcript and analysis"""
        try:
            doc = Document()

            # Set document styling
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)

            # Title
            title = doc.add_heading('Транскрипция аудио/видео', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            doc.add_paragraph(f'Дата создания: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
            doc.add_paragraph(f'ID задачи: {task_id}')
            doc.add_paragraph('')

            # Summary section
            if analysis.get('summary'):
                doc.add_heading('Краткое резюме', level=1)
                doc.add_paragraph(analysis['summary'])
                doc.add_paragraph('')

            # Key points section
            if analysis.get('key_points'):
                doc.add_heading('Ключевые моменты', level=1)
                for point in analysis['key_points']:
                    doc.add_paragraph(point, style='List Bullet')
                doc.add_paragraph('')

            # Action items section
            if analysis.get('action_items'):
                doc.add_heading('Действия и задачи', level=1)
                for action in analysis['action_items']:
                    doc.add_paragraph(action, style='List Bullet')
                doc.add_paragraph('')

            # Full transcript section
            doc.add_heading('Полная транскрипция', level=1)

            # Split transcript into paragraphs for better readability
            paragraphs = transcript.split('. ')
            for i in range(0, len(paragraphs), 5):
                chunk = '. '.join(paragraphs[i:i+5])
                if chunk:
                    doc.add_paragraph(chunk.strip() + '.')

            # Save document
            output_path = self.upload_dir / f"{task_id}_transcript.docx"
            doc.save(str(output_path))

            logger.info(f"docx_created output_path={str(output_path)}")
            return output_path

        except Exception as e:
            logger.error(f"docx_creation_error: {e}")
            raise

    async def _convert_to_mp3(self, input_path: Path, output_path: Path) -> None:
        """Convert audio to MP3 format"""
        try:
            if input_path.suffix.lower() == '.mp3':
                # Already MP3, just copy
                async with aiofiles.open(input_path, 'rb') as src:
                    content = await src.read()
                async with aiofiles.open(output_path, 'wb') as dst:
                    await dst.write(content)
                return

            # Convert using ffmpeg
            process = await asyncio.create_subprocess_exec(
                'ffmpeg',
                '-i', str(input_path),
                '-acodec', 'libmp3lame',
                '-q:a', '2',
                str(output_path),
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )

            stdout, stderr = await process.communicate()

            if process.returncode != 0:
                raise Exception(f"FFmpeg conversion error: {stderr.decode()}")

            logger.info(f"audio_converted output_path={str(output_path)}")

        except Exception as e:
            logger.error(f"audio_conversion_error: {e}")
            raise

    async def save_uploaded_file(self, file_content: bytes, filename: str) -> Path:
        """Save uploaded file to disk"""
        try:
            task_id = str(uuid.uuid4())
            file_extension = Path(filename).suffix
            file_path = self.upload_dir / f"{task_id}{file_extension}"

            async with aiofiles.open(file_path, "wb") as f:
                await f.write(file_content)

            logger.info(f"file_saved file_path={file_path}, size={len(file_content)}")

            # АВТООЧИСТКА: запускаем очистку старых файлов в фоне
            # Удаляем файлы старше 7 дней для экономии места
            asyncio.create_task(self.cleanup_old_files(days_old=7))

            return file_path

        except Exception as e:
            logger.error(f"file_save_error: {e}")
            raise

    def get_task_status(self, task_id: str) -> Optional[Dict]:
        """Get task status by ID"""
        return self.tasks.get(task_id)

    def get_file_path(self, filename: str) -> Optional[Path]:
        """Get file path for download"""
        file_path = self.upload_dir / filename
        if file_path.exists():
            return file_path
        return None

    async def combine_chunks_and_process(
        self,
        session_id: str,
        filename: str,
        task_id: str
    ) -> None:
        """Combine uploaded chunks and start transcription processing"""
        try:
            import aiofiles

            # Update task status
            self.tasks[task_id]["stage"] = "Объединение файлов..."
            self.tasks[task_id]["progress"] = 5

            session_dir = self.upload_dir / session_id

            if not session_dir.exists():
                raise Exception(f"Session directory not found: {session_id}")

            # Get all chunks sorted by index
            chunks = sorted(
                session_dir.glob("chunk_*"),
                key=lambda x: int(x.name.split("_")[1])
            )

            if not chunks:
                raise Exception("No chunks found")

            logger.info(
                "combining_chunks",
                task_id=task_id,
                session_id=session_id,
                chunk_count=len(chunks)
            )

            # Combine chunks into final file
            file_extension = Path(filename).suffix
            final_filename = f"upload_{task_id}{file_extension}"
            final_path = self.upload_dir / final_filename

            async with aiofiles.open(final_path, "wb") as outfile:
                for i, chunk_path in enumerate(chunks):
                    async with aiofiles.open(chunk_path, "rb") as infile:
                        content = await infile.read()
                        await outfile.write(content)

                    # Update progress during combination
                    progress = 5 + int((i + 1) / len(chunks) * 10)
                    self.tasks[task_id]["progress"] = progress

            # Delete chunks and session directory
            for chunk_path in chunks:
                chunk_path.unlink()

            try:
                session_dir.rmdir()
            except:
                pass

            logger.info(
                "chunks_combined",
                task_id=task_id,
                final_path=str(final_path),
                file_size=final_path.stat().st_size
            )

            # Update task before processing
            self.tasks[task_id]["stage"] = "Начало обработки видео..."
            self.tasks[task_id]["progress"] = 15

            # Now process the combined file
            await self.process_audio(final_path, task_id)

        except Exception as e:
            logger.error(
                "chunk_combination_error",
                task_id=task_id,
                error=str(e),
                exc_info=True
            )
            self.tasks[task_id] = {
                "status": "error",
                "progress": 0,
                "error": f"Ошибка при объединении файлов: {e}",
                "failed_at": datetime.now().isoformat()
            }


# Global service instance
transcription_service = TranscriptionService()
