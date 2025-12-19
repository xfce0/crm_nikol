import os
import io
import tempfile
import aiofiles
from typing import Optional, Dict, Any
from telegram import Document, Bot
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument

from ..config.settings import settings
from ..config.logging import get_logger, log_api_call

logger = get_logger(__name__)

class FileService:
    """Сервис для обработки файлов"""
    
    def __init__(self):
        self.allowed_extensions = {
            '.pdf': self._extract_pdf_text,
            '.doc': self._extract_doc_text,
            '.docx': self._extract_docx_text,
            '.txt': self._extract_txt_text
        }
        self.max_file_size = settings.MAX_FILE_SIZE
    
    async def process_document(self, document: Document, bot: Bot) -> Optional[str]:
        """Обработка загруженного документа"""
        try:
            # Проверяем размер файла
            if document.file_size > self.max_file_size:
                logger.warning(f"Файл слишком большой: {document.file_size} bytes")
                return None
            
            # Получаем расширение файла
            file_extension = self._get_file_extension(document.file_name)
            
            if file_extension not in self.allowed_extensions:
                logger.warning(f"Неподдерживаемый формат файла: {file_extension}")
                return None
            
            # Скачиваем файл
            file_content = await self._download_file(document, bot)
            
            if not file_content:
                return None
            
            # Извлекаем текст в зависимости от типа файла
            extractor = self.allowed_extensions[file_extension]
            extracted_text = await extractor(file_content, document.file_name)
            
            if extracted_text:
                logger.info(f"Успешно извлечен текст из {document.file_name}: {len(extracted_text)} символов")
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"Ошибка при обработке документа {document.file_name}: {e}")
            return None
    
    async def _download_file(self, document: Document, bot: Bot) -> Optional[bytes]:
        """Скачивание файла из Telegram"""
        try:
            # Получаем файл
            file = await bot.get_file(document.file_id)
            
            # Скачиваем содержимое файла
            file_content = await file.download_as_bytearray()
            
            log_api_call("Telegram", "download_file", True)
            
            return bytes(file_content)
            
        except Exception as e:
            log_api_call("Telegram", "download_file", False)
            logger.error(f"Ошибка при скачивании файла: {e}")
            return None
    
    def _get_file_extension(self, filename: str) -> str:
        """Получение расширения файла"""
        return '.' + filename.split('.')[-1].lower() if '.' in filename else ''
    
    async def _extract_pdf_text(self, file_content: bytes, filename: str) -> Optional[str]:
        """Извлечение текста из PDF файла"""
        try:
            text_parts = []
            
            # Создаем временный файл
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Пробуем pdfplumber для лучшего извлечения текста
                with pdfplumber.open(temp_file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                
                # Если pdfplumber не сработал, используем PyPDF2
                if not text_parts:
                    with open(temp_file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        for page in pdf_reader.pages:
                            text = page.extract_text()
                            if text:
                                text_parts.append(text)
                
            finally:
                # Удаляем временный файл
                os.unlink(temp_file_path)
            
            extracted_text = '\n\n'.join(text_parts).strip()
            
            if not extracted_text:
                logger.warning(f"Не удалось извлечь текст из PDF: {filename}")
                return None
            
            return self._clean_text(extracted_text)
            
        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из PDF {filename}: {e}")
            return None
    
    async def _extract_docx_text(self, file_content: bytes, filename: str) -> Optional[str]:
        """Извлечение текста из DOCX файла"""
        try:
            # Создаем временный файл
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Открываем документ
                doc = DocxDocument(temp_file_path)
                
                # Извлекаем текст из параграфов
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text.strip())
                
                # Извлекаем текст из таблиц
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text.strip())
                
            finally:
                # Удаляем временный файл
                os.unlink(temp_file_path)
            
            extracted_text = '\n\n'.join(text_parts).strip()
            
            if not extracted_text:
                logger.warning(f"Не удалось извлечь текст из DOCX: {filename}")
                return None
            
            return self._clean_text(extracted_text)
            
        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из DOCX {filename}: {e}")
            return None
    
    async def _extract_doc_text(self, file_content: bytes, filename: str) -> Optional[str]:
        """Извлечение текста из DOC файла (упрощенная обработка)"""
        try:
            # Для DOC файлов используем базовое извлечение
            # В продакшене лучше использовать python-docx2txt или antiword
            
            text = file_content.decode('utf-8', errors='ignore')
            
            # Простая очистка от служебных символов
            cleaned_text = ''.join(char for char in text if char.isprintable() or char.isspace())
            
            if len(cleaned_text.strip()) < 50:
                logger.warning(f"Слишком мало текста извлечено из DOC: {filename}")
                return None
            
            return self._clean_text(cleaned_text)
            
        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из DOC {filename}: {e}")
            return None
    
    async def _extract_txt_text(self, file_content: bytes, filename: str) -> Optional[str]:
        """Извлечение текста из TXT файла"""
        try:
            # Пробуем разные кодировки
            encodings = ['utf-8', 'windows-1251', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return self._clean_text(text)
                except UnicodeDecodeError:
                    continue
            
            # Если ни одна кодировка не подошла
            logger.warning(f"Не удалось декодировать TXT файл: {filename}")
            return None
            
        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из TXT {filename}: {e}")
            return None
    
    def _clean_text(self, text: str) -> str:
        """Очистка извлеченного текста"""
        try:
            # Удаляем лишние пробелы и переносы строк
            lines = text.split('\n')
            cleaned_lines = []
            
            for line in lines:
                line = line.strip()
                if line and len(line) > 2:  # Игнорируем очень короткие строки
                    cleaned_lines.append(line)
            
            # Объединяем строки
            cleaned_text = '\n'.join(cleaned_lines)
            
            # Удаляем повторяющиеся переносы строк
            import re
            cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
            
            return cleaned_text.strip()
            
        except Exception as e:
            logger.error(f"Ошибка при очистке текста: {e}")
            return text  # Возвращаем оригинальный текст если очистка не удалась
    
    async def save_file(self, file_content: bytes, filename: str, file_type: str = "document") -> Optional[str]:
        """Сохранение файла на диск"""
        try:
            # Определяем папку для сохранения
            folder_map = {
                "document": settings.DOCUMENTS_DIR,
                "image": settings.IMAGES_DIR,
                "audio": settings.AUDIO_DIR,
                "temp": settings.TEMP_DIR
            }
            
            folder = folder_map.get(file_type, settings.TEMP_DIR)
            
            # Создаем папку если не существует
            os.makedirs(folder, exist_ok=True)
            
            # Генерируем уникальное имя файла
            import uuid
            import time
            
            file_extension = self._get_file_extension(filename)
            unique_filename = f"{int(time.time())}_{uuid.uuid4().hex[:8]}{file_extension}"
            file_path = os.path.join(folder, unique_filename)
            
            # Сохраняем файл
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(file_content)
            
            logger.info(f"Файл сохранен: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Ошибка при сохранении файла {filename}: {e}")
            return None
    
    async def delete_file(self, file_path: str) -> bool:
        """Удаление файла"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Файл удален: {file_path}")
                return True
            else:
                logger.warning(f"Файл не найден для удаления: {file_path}")
                return False
                
        except Exception as e:
            logger.error(f"Ошибка при удалении файла {file_path}: {e}")
            return False
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Получение информации о файле"""
        try:
            if not os.path.exists(file_path):
                return {}
            
            stat = os.stat(file_path)
            
            return {
                "size": stat.st_size,
                "created": stat.st_ctime,
                "modified": stat.st_mtime,
                "extension": self._get_file_extension(file_path),
                "name": os.path.basename(file_path)
            }
            
        except Exception as e:
            logger.error(f"Ошибка при получении информации о файле {file_path}: {e}")
            return {}
    
    def cleanup_temp_files(self, max_age_hours: int = 24) -> int:
        """Очистка временных файлов старше указанного времени"""
        try:
            import time
            
            temp_dir = settings.TEMP_DIR
            if not os.path.exists(temp_dir):
                return 0
            
            current_time = time.time()
            max_age_seconds = max_age_hours * 3600
            deleted_count = 0
            
            for filename in os.listdir(temp_dir):
                file_path = os.path.join(temp_dir, filename)
                
                if os.path.isfile(file_path):
                    file_age = current_time - os.path.getctime(file_path)
                    
                    if file_age > max_age_seconds:
                        try:
                            os.remove(file_path)
                            deleted_count += 1
                        except Exception as e:
                            logger.error(f"Не удалось удалить временный файл {file_path}: {e}")
            
            if deleted_count > 0:
                logger.info(f"Удалено {deleted_count} временных файлов")
            
            return deleted_count
            
        except Exception as e:
            logger.error(f"Ошибка при очистке временных файлов: {e}")
            return 0

# Создаем глобальный экземпляр сервиса
file_service = FileService()

# Функция-обертка для удобства использования
async def process_uploaded_file(document: Document, bot: Bot) -> Optional[str]:
    """Обработка загруженного файла (функция-обертка)"""
    return await file_service.process_document(document, bot)